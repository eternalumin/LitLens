from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import Dict, Any, Optional
import io
import base64
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
import uuid
import os

router = APIRouter()

class ExportRequest(BaseModel):
    summary_data: Dict[Any, Any]
    format: str  # "pdf" or "docx"

@router.post("/pdf")
async def export_to_pdf(request: ExportRequest):
    """
    Export summary data to PDF format
    """
    try:
        # Create a buffer to hold the PDF
        buffer = io.BytesIO()
        
        # Create the PDF object
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        story.append(Paragraph(f"LitLens Analysis: {request.summary_data.get('title', 'Unknown Book')}", title_style))
        story.append(Spacer(1, 12))
        
        # Author
        if request.summary_data.get('author'):
            author_style = ParagraphStyle(
                'Author',
                parent=styles['Normal'],
                fontSize=16,
                spaceAfter=20,
                alignment=1
            )
            story.append(Paragraph(f"by {request.summary_data['author']}", author_style))
            story.append(Spacer(1, 12))
        
        # Overview
        story.append(Paragraph("Overview", styles['Heading2']))
        story.append(Paragraph(request.summary_data.get('overview', 'No overview available'), styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Themes
        if request.summary_data.get('themes'):
            story.append(Paragraph("Key Themes", styles['Heading2']))
            themes_text = ", ".join(request.summary_data['themes'])
            story.append(Paragraph(themes_text, styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Characters
        if request.summary_data.get('characters'):
            story.append(Paragraph("Main Characters", styles['Heading2']))
            for char in request.summary_data['characters']:
                char_text = f"<b>{char.get('name', 'Unknown')}</b>: {char.get('description', 'No description available')}"
                story.append(Paragraph(char_text, styles['Normal']))
                story.append(Spacer(1, 6))
            story.append(Spacer(1, 12))
        
        # Quotes
        if request.summary_data.get('quotes'):
            story.append(Paragraph("Important Quotes", styles['Heading2']))
            for quote in request.summary_data['quotes'][:5]:  # Limit to 5 quotes
                quote_text = f'"{quote.get("quote", "")}"'
                context_text = f"<i>{quote.get('context', '')}</i>"
                story.append(Paragraph(quote_text, styles['Normal']))
                story.append(Paragraph(context_text, styles['Italic']))
                story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        
        # Get the value of the BytesIO buffer
        pdf_value = buffer.getvalue()
        buffer.close()
        
        # Encode to base64 for transport
        pdf_base64 = base64.b64encode(pdf_value).decode('utf-8')
        
        return {
            "filename": f"LitLens_Analysis_{request.summary_data.get('title', 'book').replace(' ', '_')}.pdf",
            "content": pdf_base64,
            "content_type": "application/pdf"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating PDF: {str(e)}")

@router.post("/docx")
async def export_to_docx(request: ExportRequest):
    """
    Export summary data to DOCX format
    """
    try:
        # Create a new Document
        doc = Document()
        
        # Title
        title = doc.add_heading(f'LitLens Analysis: {request.summary_data.get("title", "Unknown Book")}', 0)
        title.alignment = 1  # Center alignment
        
        # Author
        if request.summary_data.get('author'):
            author_para = doc.add_paragraph(f'by {request.summary_data["author"]}')
            author_para.alignment = 1
            author_para.runs[0].italic = True
        
        doc.add_paragraph()  # Empty line
        
        # Overview
        doc.add_heading('Overview', level=1)
        doc.add_paragraph(request.summary_data.get('overview', 'No overview available'))
        
        # Themes
        if request.summary_data.get('themes'):
            doc.add_heading('Key Themes', level=1)
            themes_para = doc.add_paragraph(', '.join(request.summary_data['themes']))
        
        # Characters
        if request.summary_data.get('characters'):
            doc.add_heading('Main Characters', level=1)
            for char in request.summary_data['characters']:
                char_para = doc.add_paragraph()
                char_para.add_run(f"{char.get('name', 'Unknown')}: ").bold = True
                char_para.add_run(char.get('description', 'No description available'))
        
        # Quotes
        if request.summary_data.get('quotes'):
            doc.add_heading('Important Quotes', level=1)
            for quote in request.summary_data['quotes'][:5]:  # Limit to 5 quotes
                quote_para = doc.add_paragraph()
                quote_para.add_run('"').italic = True
                quote_para.add_run(quote.get('quote', ''))
                quote_para.add_run('"').italic = True
                quote_para.add_run(f" — {quote.get('context', '')}").italic = True
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Encode to base64
        docx_value = buffer.getvalue()
        buffer.close()
        docx_base64 = base64.b64encode(docx_value).decode('utf-8')
        
        return {
            "filename": f"LitLens_Analysis_{request.summary_data.get('title', 'book').replace(' ', '_')}.docx",
            "content": docx_base64,
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")
